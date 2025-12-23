import streamlit as st
import re
import fitz  # PyMuPDF
import docx
import os
import json
from datetime import datetime
import csv
import tempfile
import pandas as pd

def read_pdf(file_path):
    try:
        pdf = fitz.open(file_path)
        text = ""
        for page in pdf:
            text += page.get_text("text") + "\n"
        pdf.close()
        return text
    except Exception as e:
        return f"PDF error: {str(e)}"




def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                text += " " + " ".join(cell.text for cell in row.cells)
        return text
    except Exception as e:
        return f"DOCX error: {str(e)}"




def read_resume(file_path):
    if not os.path.exists(file_path):
        return "File not found!"
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext == ".docx":
        return read_docx(file_path)
    return "Unsupported file type!"




# ---------- Cleaning ----------
def clean_text(text: str) -> str:
    """Aggressive cleaning for skills, email, phone extraction"""
    if not isinstance(text, str):
        return ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()




def clean_text_keep_lines(text: str) -> str:
    """Clean text but preserve line structure for section detection"""
    if not isinstance(text, str):
        return ""
    lines = text.split("\n")
    cleaned_lines = []
    for ln in lines:
        ln = re.sub(r"\s+", " ", ln).strip()
        if ln:
            cleaned_lines.append(ln)
    return "\n".join(cleaned_lines)




# ---------- Basic info extraction ----------
EMAIL_REGEX = re.compile(
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
)



PHONE_REGEX = re.compile(
    r"(?:\+?91[\s\-]*)?[6-9]\d{9}"
)



def extract_basic_info(text: str) -> dict:
    info = {}

    # Emails
    emails = EMAIL_REGEX.findall(text)
    info["email"] = emails[0] if emails else "Not found"
    info["emails"] = list(dict.fromkeys(emails))

    # Phones
    phones = PHONE_REGEX.findall(text)
    info["phone"] = phones[0] if phones else "Not found"
    info["phones"] = list(dict.fromkeys(phones))

    return info




# ---------- Skills extraction ----------
SKILLS = [
    # Programming languages
    "python", "java", "c", "c++", "c#", "javascript", "typescript", "php",
    "go", "ruby", "kotlin", "swift",
    # Web
    "html", "css", "react", "angular", "vue", "node.js", "django", "flask",
    # Data / ML
    "sql", "mysql", "postgresql", "mongodb", "pandas", "numpy",
    "machine learning", "deep learning", "data analysis", "data science",
    "tensorflow", "pytorch", "scikit-learn",
    # Tools
    "excel", "git", "docker", "kubernetes", "linux"
]



def build_skill_patterns(skills):
    multi = []
    single = []
    for s in skills:
        if " " in s:  # multi-word skill
            pattern = re.compile(r"\b" + re.escape(s) + r"\b", re.IGNORECASE)
            multi.append((s, pattern))
        else:
            pattern = re.compile(r"\b" + re.escape(s) + r"\b", re.IGNORECASE)
            single.append((s, pattern))
    return multi, single



MULTIWORD_SKILLS, SINGLEWORD_SKILLS = build_skill_patterns(SKILLS)



def extract_skills(text: str):
    if not isinstance(text, str):
        return []

    found = set()

    # 1) Check multi-word skills first
    for skill, pattern in MULTIWORD_SKILLS:
        if pattern.search(text):
            found.add(skill)

    # 2) Then check single-word skills
    for skill, pattern in SINGLEWORD_SKILLS:
        if pattern.search(text):
            found.add(skill)

    return sorted(found, key=str.lower)




# ---------- Education & Experience extraction ----------
EDU_KEYWORDS = [
    "education", "academic", "degree", "degrees", "college",
    "university", "b.tech", "btech", "b.e", "be", "bachelor",
    "master", "m.tech", "mtech", "bsc", "msc", "diploma"
]



EXP_KEYWORDS = [
    "experience", "work experience", "professional experience",
    "employment", "internship", "intern", "company",
    "organization", "worked", "job", "role", "position"
]



def extract_sections(text: str):
    """
    Returns:
        education_lines: list of lines likely related to education
        experience_lines: list of lines likely related to experience
    """
    if not isinstance(text, str):
        return [], []

    # Work with line structure (better for resumes)
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]

    education_lines = []
    experience_lines = []

    for line in lines:
        lower = line.lower()
        if len(lower) < 3:
            continue

        if any(kw in lower for kw in EDU_KEYWORDS):
            education_lines.append(line)
            continue

        if any(kw in lower for kw in EXP_KEYWORDS):
            experience_lines.append(line)

    def dedup(seq):
        seen = set()
        out = []
        for x in seq:
            if x not in seen:
                seen.add(x)
                out.append(x)
        return out

    return dedup(education_lines), dedup(experience_lines)


# ---------- MAIN PARSER ----------
def parse_resume(file_path: str) -> dict:
    """
    Complete resume parser: reads file -> extracts all info -> returns structured data
    """
    # Step 1: Read raw text
    raw_text = read_resume(file_path)
    if raw_text.startswith(("PDF error:", "DOCX error:", "File not found!", "Unsupported")):
        return {"error": raw_text, "parsed_data": {}}
    
    # Step 2: Two cleaning strategies
    cleaned_text = clean_text(raw_text)                         # Aggressive: for skills, email, phone
    cleaned_for_sections = clean_text_keep_lines(raw_text)      # Line-aware: for education/experience
    
    # Step 3: Extract everything
    basic_info = extract_basic_info(cleaned_text)
    skills = extract_skills(cleaned_text)
    education, experience = extract_sections(cleaned_for_sections)
    
    # Step 4: Structured output with metadata
    parsed_data = {
        "contact": {
            "email": basic_info["email"],
            "phone": basic_info["phone"],
            "all_emails": basic_info["emails"],
            "all_phones": basic_info["phones"]
        },
        "skills": skills,
        "education": education,
        "experience": experience,
        "summary": {
            "total_skills_found": len(skills),
            "education_lines": len(education),
            "experience_lines": len(experience),
            "raw_text_length": len(raw_text)
        }
    }
    
    return parsed_data

def save_to_csv(results, filename="resume_results.csv"):
    """
    Save resume parsing results to CSV with flattened nested data.
    Handles errors, different structures, and creates timestamped files.
    """
    if not results:
        print("⚠️  No results to save!")
        return None
    
    print(f"💾 Saving {len(results)} resumes to CSV...")
    
    # Filter successful parses only
    valid_results = [r for r in results if "error" not in r]
    failed_count = len(results) - len(valid_results)
    
    if not valid_results:
        print("❌ No valid resumes to export!")
        return None
    
    # Define ALL possible columns (flattened structure)
    fieldnames = [
        'file_name', 'file_size_kb', 'email', 'phone', 
        'total_skills', 'skills_list', 'education_lines', 'experience_lines',
        'raw_text_length'
    ]
    
    # Prepare flattened rows
    csv_rows = []
    for resume in valid_results:
        file_info = resume.get("file_info", {})
        contact = resume.get("contact", {})
        summary = resume.get("summary", {})
        skills = resume.get("skills", [])
        
        row = {
            'file_name': file_info.get('name', 'Unknown'),
            'file_size_kb': file_info.get('size_kb', 0),
            'email': contact.get('email', 'Not found'),
            'phone': contact.get('phone', 'Not found'),
            'total_skills': summary.get('total_skills_found', 0),
            'skills_list': ', '.join(skills[:5]) + ('...' if len(skills) > 5 else ''),
            'education_lines': summary.get('education_lines', 0),
            'experience_lines': summary.get('experience_lines', 0),
            'raw_text_length': summary.get('raw_text_length', 0)
        }
        csv_rows.append(row)
    
    # Create timestamped filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = filename.replace('.csv', f'_{timestamp}.csv')
    
    try:
        with open(safe_filename, "w", newline="", encoding="utf-8") as file:
            writer = csv.DictWriter(file, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(csv_rows)
        
        print(f"✅ CSV saved: {safe_filename}")
        print(f"📊 Exported: {len(csv_rows)} valid / {failed_count} failed resumes")
        return safe_filename
        
    except Exception as e:
        print(f"❌ CSV save failed: {str(e)}")
        return None


def save_search_to_csv(search_results, filename="skill_search_results.csv"):
    """Save skill search results to CSV"""
    if not search_results:
        print("⚠️  No search results to save!")
        return None
    
    fieldnames = ['file_name', 'file_size_kb', 'matched_skills', 'match_count', 'total_skills']
    
    csv_rows = []
    for result in search_results:
        row = {
            'file_name': result.get('name', 'Unknown'),
            'file_size_kb': result.get('size_kb', 0),
            'matched_skills': ', '.join([m['skill'] for m in result.get('matched_skills', [])]),
            'match_count': result.get('match_count', 0),
            'total_skills': result.get('total_skills', 0)
        }
        csv_rows.append(row)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = filename.replace('.csv', f'_{timestamp}.csv')
    
    with open(safe_filename, "w", newline="", encoding="utf-8") as file:
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(csv_rows)
    
    print(f"✅ Search CSV saved: {safe_filename}")
    return safe_filename



# ========== NEW SKILL SEARCH FUNCTIONS ==========
def search_by_skill(resumes, skill, case_sensitive=False, partial_match=True):
    """
    Advanced resume search by skill with multiple matching strategies.
    
    Args:
        resumes: List of parsed resume dicts
        skill: Skill to search for (str)
        case_sensitive: If True, respect case (default: False)
        partial_match: If True, match partial skills (default: True)
    
    Returns:
        List of matching resumes with match details
    """
    if not skill or not resumes:
        return []
    
    skill = skill if case_sensitive else skill.lower().strip()
    matched = []
    
    for resume in resumes:
        if "error" in resume:  # Skip failed parses
            continue
            
        resume_skills = resume.get("skills", [])
        matches = []
        
        for resume_skill in resume_skills:
            resume_skill_clean = resume_skill if case_sensitive else resume_skill.lower()
            
            if partial_match:
                # Flexible partial matching (e.g., "py" matches "python")
                if skill in resume_skill_clean or resume_skill_clean in skill:
                    matches.append({
                        "skill": resume_skill,
                        "match_type": "partial" if len(skill) < len(resume_skill_clean) else "exact"
                    })
            else:
                # Exact word boundary matching only
                if skill == resume_skill_clean:
                    matches.append({
                        "skill": resume_skill,
                        "match_type": "exact"
                    })
        
        if matches:
            matched.append({
                **resume["file_info"],  # File metadata
                "matched_skills": matches,
                "total_skills": len(resume_skills),
                "match_count": len(matches)
            })
    
    # Sort by most matches first
    matched.sort(key=lambda x: x["match_count"], reverse=True)
    return matched

def main():
    st.set_page_config(page_title="Smart Resume Parser", page_icon="🧠", layout="wide")
    
    st.title("💡 Smart Resume Parser")
    st.markdown("---")
    
    # Sidebar
    st.sidebar.title("💡 Smart Resume Parser")
    st.sidebar.markdown("---")
    st.sidebar.header("⚙️ Options")
    max_files = st.sidebar.slider("Max files to process", 1, 50, 10)
    export_csv = st.sidebar.checkbox("Export to CSV", value=True)
    
    # Main tabs
    tab1, tab2, tab3 = st.tabs(["📁 Upload & Parse", "🔍 Skill Search", "📊 Analytics"])
    
    with tab1:
        st.header("Upload Resumes")
        
        uploaded_files = st.file_uploader(
            "Choose PDF/DOCX files",
            type=["pdf", "docx"],
            accept_multiple_files=True,
            help="Supports multiple files (max 10MB each)"
        )
        
        if uploaded_files:
            st.success(f"✅ {len(uploaded_files)} files uploaded!")
            
            if st.button("🚀 Parse All Resumes", type="primary"):
                with st.spinner("Parsing resumes..."):
                    results = []
                    
                    # Process each file
                    for i, uploaded_file in enumerate(uploaded_files[:max_files]):
                        # Save uploaded file to temp
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f"_{uploaded_file.name}") as tmp:
                            tmp.write(uploaded_file.getvalue())
                            tmp_path = tmp.name
                        
                        st.info(f"📄 Processing {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
                        
                        # Parse with your function
                        result = parse_resume(tmp_path)
                        
                        # Add file info
                        result["file_info"] = {
                            "name": uploaded_file.name,
                            "size_kb": len(uploaded_file.getvalue()) / 1024
                        }
                        
                        results.append(result)
                        os.unlink(tmp_path)  # Clean up temp file
                    
                    st.session_state.results = results
                    st.success("🎉 Parsing complete!")
                
                # Display results
                if 'results' in st.session_state:
                    display_parse_results(st.session_state.results)
    
    with tab2:
        st.header("🔍 Skill Search")
        if 'results' in st.session_state:
            skill = st.text_input("Enter skill to search (e.g., 'python', 'react')")
            if skill:
                if st.button("Search", type="secondary"):
                    matches = search_by_skill(st.session_state.results, skill)
                    display_search_results(matches, skill)
        else:
            st.warning("👈 Parse some resumes first!")
    
    with tab3:
        st.header("📊 Analytics")
        if 'results' in st.session_state:
            display_analytics(st.session_state.results, export_csv)
        else:
            st.warning("👈 Parse some resumes first!")

def display_parse_results(results):
    st.subheader(f"📋 Results: {len([r for r in results if 'error' not in r])} successful / {len(results)} total")
    
    for result in results:
        if "error" in result:
            st.error(f"❌ {result['file_info']['name']}: {result['error'][:100]}...")
        else:
            with st.expander(f"✅ {result['file_info']['name']} ({result['summary']['total_skills_found']} skills)"):
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Email", result['contact']['email'])
                    st.metric("Phone", result['contact']['phone'])
                    st.metric("Skills", len(result['skills']))
                with col2:
                    st.metric("Education", len(result['education']))
                    st.metric("Experience", len(result['experience']))
                    st.write("**Skills:**", ", ".join(result['skills'][:5]))
    
    if st.button("💾 Save JSON"):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"streamlit_results_{timestamp}.json"
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(st.session_state.results, f, indent=2, ensure_ascii=False)
        st.success(f"💾 Saved: {filename}")
        st.download_button("📥 Download JSON", data=open(filename, 'rb').read(), file_name=filename)

def display_search_results(matches, skill):
    st.success(f"🎯 Found {len(matches)} resumes with '{skill}'")
    
    for i, match in enumerate(matches):
        with st.expander(f"#{i+1} {match['name']} ({match['match_count']} matches)"):
            st.write(f"**Total Skills:** {match['total_skills']}")
            for m in match['matched_skills']:
                st.success(f"✅ {m['skill']} ({m['match_type'].upper()})")

def display_analytics(results, export_csv):
    valid_results = [r for r in results if "error" not in r]
    
    if valid_results:
        df = pd.DataFrame([{
            'File': r['file_info']['name'],
            'Email': r['contact']['email'],
            'Skills': len(r['skills']),
            'Education': len(r['education']),
            'Experience': len(r['experience'])
        } for r in valid_results])
        
        st.dataframe(df, use_container_width=True)
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Skills", sum(len(r['skills']) for r in valid_results))
        with col2:
            st.metric("Avg Skills", df['Skills'].mean())
        with col3:
            st.metric("Success Rate", f"{len(valid_results)}/{len(results)}")
        
        if export_csv:
            csv = df.to_csv(index=False).encode('utf-8')
            st.download_button("📊 Download CSV", csv, "analytics.csv", "text/csv")

if __name__ == "__main__":
    main()
